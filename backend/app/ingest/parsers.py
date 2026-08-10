"""Safe TXT/Markdown/DOCX/PDF/FDX parsing with source provenance."""

from __future__ import annotations

import hashlib
import io
import re
import stat
import uuid
import zipfile
from pathlib import PurePosixPath

from lxml import etree

from app.schema.studio import SourceDocument, SourceSpan


class SourceIngestError(ValueError):
    pass


class SourceIngestor:
    MAX_BYTES = 25 * 1024 * 1024
    MAX_ARCHIVE_ENTRIES = 2_000
    MAX_ARCHIVE_UNCOMPRESSED = 100 * 1024 * 1024
    MAX_PDF_PAGES = 1_000

    def ingest(self, filename: str, content: bytes) -> SourceDocument:
        if not filename or len(filename) > 255:
            raise SourceIngestError("invalid source filename")
        if not content:
            raise SourceIngestError("source is empty")
        if len(content) > self.MAX_BYTES:
            raise SourceIngestError("source exceeds 25 MB")
        suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        if suffix in {"txt"}:
            text, source_format = self._decode_text(content), "text"
        elif suffix in {"md", "markdown"}:
            text, source_format = self._decode_text(content), "markdown"
        elif suffix == "docx":
            text, source_format = self._docx(content), "docx"
        elif suffix == "pdf":
            text, source_format = self._pdf(content), "pdf"
        elif suffix == "fdx":
            text, source_format = self._fdx(content), "fdx"
        else:
            raise SourceIngestError("supported source formats are TXT, Markdown, DOCX, PDF and FDX")
        normalized = self._normalize(text)
        if not normalized:
            raise SourceIngestError("source contains no extractable text")
        source_id = f"src_{uuid.uuid4().hex}"
        return SourceDocument(
            id=source_id,
            filename=PurePosixPath(filename.replace("\\", "/")).name,
            format=source_format,
            sha256=hashlib.sha256(content).hexdigest(),
            text=normalized,
            spans=self._spans(source_id, normalized),
        )

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise SourceIngestError("text encoding must be UTF-8 or GB18030")

    def _validate_zip(self, content: bytes) -> None:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                infos = archive.infolist()
                if len(infos) > self.MAX_ARCHIVE_ENTRIES:
                    raise SourceIngestError("document archive contains too many entries")
                if sum(info.file_size for info in infos) > self.MAX_ARCHIVE_UNCOMPRESSED:
                    raise SourceIngestError("document archive expands beyond the safe limit")
                for info in infos:
                    path = PurePosixPath(info.filename.replace("\\", "/"))
                    if path.is_absolute() or ".." in path.parts:
                        raise SourceIngestError("document archive contains an unsafe path")
                    mode = (info.external_attr >> 16) & 0xFFFF
                    if stat.S_ISLNK(mode):
                        raise SourceIngestError("document archive may not contain symlinks")
        except zipfile.BadZipFile as exc:
            raise SourceIngestError("DOCX is not a valid ZIP document") from exc

    def _docx(self, content: bytes) -> str:
        self._validate_zip(content)
        try:
            from docx import Document

            document = Document(io.BytesIO(content))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except ImportError as exc:
            raise SourceIngestError("DOCX parser is not installed; install python-docx") from exc
        except Exception as exc:
            raise SourceIngestError("DOCX could not be parsed safely") from exc

    def _pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content), strict=True)
            if reader.is_encrypted:
                raise SourceIngestError("encrypted PDF is not supported")
            if len(reader.pages) > self.MAX_PDF_PAGES:
                raise SourceIngestError("PDF contains too many pages")
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except SourceIngestError:
            raise
        except ImportError as exc:
            raise SourceIngestError("PDF parser is not installed; install pypdf") from exc
        except Exception as exc:
            raise SourceIngestError("PDF could not be parsed safely") from exc

    @staticmethod
    def _fdx(content: bytes) -> str:
        if b"<!DOCTYPE" in content.upper() or b"<!ENTITY" in content.upper():
            raise SourceIngestError("FDX external entities are forbidden")
        try:
            parser = etree.XMLParser(resolve_entities=False, no_network=True, recover=False, huge_tree=False)
            root = etree.fromstring(content, parser=parser)
            return "\n".join(
                text.strip() for text in root.xpath("//*[local-name()='Text']/text()") if text.strip()
            )
        except Exception as exc:
            raise SourceIngestError("FDX could not be parsed safely") from exc

    @staticmethod
    def _normalize(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
        text = "\n".join(line.rstrip() for line in text.split("\n"))
        return re.sub(r"\n{3,}", "\n\n", text).strip()

    @staticmethod
    def _spans(source_id: str, text: str) -> list[SourceSpan]:
        spans: list[SourceSpan] = []
        for index, match in enumerate(re.finditer(r"[^\n]+(?:\n(?!\n)[^\n]+)*", text), start=1):
            value = match.group(0).strip()
            if not value:
                continue
            start = match.start() + len(match.group(0)) - len(match.group(0).lstrip())
            end = start + len(value)
            line_start = text.count("\n", 0, start) + 1
            line_end = text.count("\n", 0, end) + 1
            spans.append(SourceSpan(
                id=f"{source_id}_span_{index}", source_id=source_id,
                start=start, end=end, line_start=line_start, line_end=line_end, text=value,
            ))
        return spans
